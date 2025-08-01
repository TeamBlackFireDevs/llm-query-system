import os
import hashlib
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import aiofiles
import PyPDF2
from docx import Document as DocxDocument
import email
from email.policy import default
from sqlalchemy.orm import Session
from models.database_models import Document, DocumentChunk
from models.pydantic_models import DocumentType
from config import settings

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)
        
    async def process_document(self, file_content: bytes, filename: str, content_type: str, db: Session) -> Document:
        """Process uploaded document and store in database"""
        try:
            # Determine document type
            doc_type = self._get_document_type(filename, content_type)
            
            # Generate unique filename
            file_hash = hashlib.md5(file_content).hexdigest()
            unique_filename = f"{file_hash}_{filename}"
            file_path = self.upload_dir / unique_filename
            
            # Save file
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            # Create document record
            document = Document(
                filename=unique_filename,
                original_filename=filename,
                file_path=str(file_path),
                file_size=len(file_content),
                content_type=content_type,
                document_type=doc_type.value,
                processing_status="processing"
            )
            
            db.add(document)
            db.commit()
            db.refresh(document)
            
            # Extract text and create chunks
            try:
                text_content = await self._extract_text(file_path, doc_type)
                chunks = self._create_chunks(text_content)
                
                # Store chunks in database
                for i, chunk_content in enumerate(chunks):
                    chunk_hash = hashlib.md5(chunk_content.encode()).hexdigest()
                    chunk = DocumentChunk(
                        document_id=document.id,
                        chunk_index=i,
                        content=chunk_content,
                        content_hash=chunk_hash,
                        metadata={
                            "chunk_size": len(chunk_content),
                            "document_type": doc_type.value
                        }
                    )
                    db.add(chunk)
                
                # Update document status
                document.processing_status = "completed"
                document.chunk_count = len(chunks)
                db.commit()
                
                logger.info(f"Successfully processed document {filename} with {len(chunks)} chunks")
                
            except Exception as e:
                document.processing_status = "failed"
                document.processing_error = str(e)
                db.commit()
                logger.error(f"Failed to process document {filename}: {str(e)}")
                raise
            
            return document
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            raise
    
    def _get_document_type(self, filename: str, content_type: str) -> DocumentType:
        """Determine document type from filename and content type"""
        extension = Path(filename).suffix.lower()
        
        if extension == '.pdf' or 'pdf' in content_type:
            return DocumentType.PDF
        elif extension == '.docx' or 'wordprocessingml' in content_type:
            return DocumentType.DOCX
        elif extension == '.txt' or 'text/plain' in content_type:
            return DocumentType.TXT
        elif extension == '.eml' or 'message/rfc822' in content_type:
            return DocumentType.EMAIL
        else:
            return DocumentType.TXT  # Default fallback
    
    async def _extract_text(self, file_path: Path, doc_type: DocumentType) -> str:
        """Extract text content from document based on type"""
        try:
            if doc_type == DocumentType.PDF:
                return await self._extract_pdf_text(file_path)
            elif doc_type == DocumentType.DOCX:
                return await self._extract_docx_text(file_path)
            elif doc_type == DocumentType.TXT:
                return await self._extract_txt_text(file_path)
            elif doc_type == DocumentType.EMAIL:
                return await self._extract_email_text(file_path)
            else:
                raise ValueError(f"Unsupported document type: {doc_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise
        return text.strip()
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise
        return text.strip()
    
    async def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                text = await file.read()
        except Exception as e:
            logger.error(f"Error extracting TXT text: {str(e)}")
            raise
        return text.strip()
    
    async def _extract_email_text(self, file_path: Path) -> str:
        """Extract text from email file"""
        try:
            with open(file_path, 'rb') as file:
                msg = email.message_from_bytes(file.read(), policy=default)
                
                text_parts = []
                
                # Extract headers
                text_parts.append(f"Subject: {msg.get('Subject', 'N/A')}")
                text_parts.append(f"From: {msg.get('From', 'N/A')}")
                text_parts.append(f"To: {msg.get('To', 'N/A')}")
                text_parts.append(f"Date: {msg.get('Date', 'N/A')}")
                text_parts.append("---")
                
                # Extract body
                if msg.is_multipart():
                    for part in msg.walk():
                        if part.get_content_type() == "text/plain":
                            text_parts.append(part.get_content())
                else:
                    if msg.get_content_type() == "text/plain":
                        text_parts.append(msg.get_content())
                
                text = "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting email text: {str(e)}")
            raise
        return text.strip()
    
    def _create_chunks(self, text: str) -> List[str]:
        """Split text into chunks for processing"""
        if not text:
            return []
        
        # Simple chunking strategy - split by sentences and group
        sentences = text.split('.')
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            # Check if adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) + 1 > settings.CHUNK_SIZE:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + "."
            else:
                current_chunk += sentence + "."
        
        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    async def delete_document(self, document_id: str, db: Session) -> bool:
        """Delete document and associated chunks"""
        try:
            # Get document
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                return False
            
            # Delete file
            file_path = Path(document.file_path)
            if file_path.exists():
                file_path.unlink()
            
            # Delete chunks
            db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
            
            # Delete document
            db.delete(document)
            db.commit()
            
            logger.info(f"Successfully deleted document {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {str(e)}")
            db.rollback()
            return False 